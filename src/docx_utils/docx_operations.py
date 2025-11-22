""" 
docx_utils/docx_operations.py 

Module for checking DOCX documents for font compliance. 
Highlights runs with incorrect fonts in red and returns the total number of discrepancies. 
"""

from typing import List, Tuple, Dict
from docx.document import Document as DocumentObject
from docx.text.run import Run
from docx import Document
from docx_utils.font_check import check_paragraph_font, check_table_font
from config.config import ReportItem
from docx_utils.alignment_check import check_alignment_and_indent
from docx_utils.page_margins import check_page_margins
from docx_utils.auto_fix.font_fix import fix_paragraph_font, fix_table_font
from docx_utils.auto_fix.page_margins_fix import fix_page_margins
from docx_utils.auto_fix.alignment_fix import fix_alignment_and_indent

def analyze_docx(docx_path: str) -> Tuple[List[ReportItem], DocumentObject]:
    """
    Checks a DOCX file for font compliance with TARGET_FONT.
    Highlights runs with incorrect fonts in red in memory only.
    
    Args:
        docx_path (str): Path to the original DOCX file.
    
    Returns:
        Tuple[List[ReportItem], DocumentObject]: A tuple containing a list of runs with font discrepancies and the loaded Document object with highlights.
    """
    docx_checked: DocumentObject = Document(docx_path)
    docx_fixed: DocumentObject = Document(docx_path)
    report: List[ReportItem] = []

    # Check all paragraphs
    for paragraph_checked, paragraph_fixed in zip(docx_checked.paragraphs, docx_fixed.paragraphs):
        # Analyze for issues
        check_paragraph_font(paragraph_checked, report)
        # Apply fixes in the fixed document
        fix_paragraph_font(paragraph_fixed)

    # Check all tables
    for table_checked, table_fixed in zip(docx_checked.tables, docx_fixed.tables):
            check_table_font(table_checked, report)
            fix_table_font(table_fixed)

    check_alignment_and_indent(docx_checked, report)
    fix_alignment_and_indent(docx_fixed)
    
    check_page_margins(docx_checked, report)
    fix_page_margins(docx_fixed)

    return report, docx_checked, docx_fixed  # Return doc object for optional saving


def save_docx(docx: DocumentObject, output_path: str):
    """
    Saves a DOCX Document object to the specified file.

    Args:
        docx (DocumentObject): The Document object to save.
        output_path (str): Path to save the DOCX file.
    """
    docx.save(output_path)
