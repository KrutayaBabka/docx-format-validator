"""
docx_utils/alignment_fix.py

Module for fixing paragraph alignment, indentation, line spacing, and captions in a DOCX document.
"""

from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.document import Document as DocumentObject
import re
from config.config import (
    FIRST_LINE_INDENT_CM,
    LEFT_INDENT_CM,
    RIGHT_INDENT_CM,
    LINE_SPACING,
    TITLE_PAGE_PATTERN
)

def is_title_page(paragraphs, index) -> bool:
    text = paragraphs[index].text.strip()
    return bool(re.search(TITLE_PAGE_PATTERN, text))

def is_image_caption(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    return bool(re.match(r"^Рис\.\s*\d*", text))

def is_table_caption(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    return bool(re.match(r"^Табл\.\s*\d*", text))

def fix_paragraph_format(paragraph: Paragraph, check_first_line: bool = True) -> None:
    """
    Fix indentation and line spacing for a paragraph.
    """
    fmt = paragraph.paragraph_format
    if check_first_line:
        fmt.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
    fmt.left_indent = Cm(LEFT_INDENT_CM)
    fmt.right_indent = Cm(RIGHT_INDENT_CM)
    fmt.line_spacing = LINE_SPACING

def fix_alignment_and_indent(docx: DocumentObject) -> None:
    """
    Fix all paragraphs and table cell paragraphs for alignment, indentation, and line spacing.
    """
    paragraphs = docx.paragraphs
    skip_until_index = -1

    # Detect title page
    for i, paragraph in enumerate(paragraphs):
        if is_title_page(paragraphs, i):
            skip_until_index = i
            break

    for i, paragraph in enumerate(paragraphs):
        if i <= skip_until_index:
            continue  # skip title page
        text = paragraph.text.strip()
        if not text:
            continue

        # Fix alignment
        if is_image_caption(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif is_table_caption(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Fix caption text to be plain (remove bold/italic/underline)
        if is_image_caption(paragraph) or is_table_caption(paragraph):
            for run in paragraph.runs:
                run.bold = False
                run.italic = False
                run.underline = False

        # Fix indentation and line spacing
        # check_first_line = not (is_image_caption(paragraph) or is_table_caption(paragraph))
        fix_paragraph_format(paragraph, check_first_line=True)

    # Fix paragraphs inside tables
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fix_paragraph_format(paragraph, check_first_line=False)
