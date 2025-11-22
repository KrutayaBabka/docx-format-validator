""" 
docx_utils/alignment_check.py

Module for checking paragraph alignment and indentation in a DOCX document.
Highlights runs with incorrect alignment in red and records discrepancies
as tuples (run, reason) in the report list. Also checks first-line indentation
for regular paragraphs (not in tables).
"""

from typing import List
from docx.shared import RGBColor, Cm
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.document import Document as DocumentObject
from config.config import (
    ReportItem, 
    TITLE_PAGE_PATTERN,
    FIRST_LINE_INDENT_CM,
    LEFT_INDENT_CM,
    RIGHT_INDENT_CM,
    LINE_SPACING
)
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH


def highlight_alignment(paragraph: Paragraph, report: List[ReportItem], reason: str) -> None:
    """
    Highlight the entire paragraph in red and append a reasoned dictionary to report.
    """
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(255, 0, 0)
    report.append({
        "run": paragraph.runs[0] if paragraph.runs else None,
        "paragraph_text": paragraph.text,
        "reason": reason
    })


def is_title_page(paragraphs: List[Paragraph], index: int) -> bool:
    """
    Determine if a paragraph belongs to the title page based on TITLE_PAGE_PATTERN.
    """
    text = paragraphs[index].text.strip()
    return bool(re.search(TITLE_PAGE_PATTERN, text))


def is_image_caption(paragraph: Paragraph) -> bool:
    """
    Check if the paragraph is a caption under an image (starts with 'Рис.').
    """
    text = paragraph.text.strip()
    return bool(re.match(r"^Рис\.\s*\d*", text))


def is_table_caption(paragraph: Paragraph) -> bool:
    """
    Check if the paragraph is a caption above a table (starts with 'Табл.').
    """
    text = paragraph.text.strip()
    return bool(re.match(r"^Табл\.\s*\d*", text))


def check_paragraph_format(paragraph: Paragraph, report: List[ReportItem], check_first_line: bool = True) -> None:
    """
    Check indentation and line spacing for a single paragraph.
    Highlights issues in red and appends them to the report.
    """
    # First-line indentation
    if check_first_line:
        actual_first_line = paragraph.paragraph_format.first_line_indent
        actual_first_cm = actual_first_line.cm if actual_first_line else 0.0
        if abs(actual_first_cm - FIRST_LINE_INDENT_CM) > 1e-2:
            highlight_alignment(paragraph, report,
                f"First-line indentation should be {FIRST_LINE_INDENT_CM} cm (found {actual_first_cm:.2f} cm)")

    # Left indent
    actual_left_indent = paragraph.paragraph_format.left_indent
    actual_left_cm = actual_left_indent.cm if actual_left_indent else 0.0
    if abs(actual_left_cm - LEFT_INDENT_CM) > 1e-2:
        highlight_alignment(paragraph, report,
            f"Left indent should be {LEFT_INDENT_CM} cm (found {actual_left_cm:.2f} cm)")

    # Right indent
    actual_right_indent = paragraph.paragraph_format.right_indent
    actual_right_cm = actual_right_indent.cm if actual_right_indent else 0.0
    if abs(actual_right_cm - RIGHT_INDENT_CM) > 1e-2:
        highlight_alignment(paragraph, report,
            f"Right indent should be {RIGHT_INDENT_CM} cm (found {actual_right_cm:.2f} cm)")

    # Line spacing
    actual_line_spacing = paragraph.paragraph_format.line_spacing
    actual_spacing = actual_line_spacing if actual_line_spacing else 1.0
    if abs(actual_spacing - LINE_SPACING) > 1e-2:
        highlight_alignment(paragraph, report,
            f"Line spacing should be {LINE_SPACING} (found {actual_spacing:.2f})")


def check_alignment_and_indent(docx: DocumentObject, report: List[ReportItem]) -> None:
    """
    Check all paragraphs and table cell paragraphs in a document for correct
    alignment, indentation, and line spacing.
    """
    paragraphs = docx.paragraphs
    skip_until_index = -1

    # Detect the title page
    for i, paragraph in enumerate(paragraphs):
        if is_title_page(paragraphs, i):
            skip_until_index = i
            break

    # Check regular paragraphs
    for i, paragraph in enumerate(paragraphs):
        if i <= skip_until_index:
            continue  # skip title page

        text = paragraph.text.strip()
        if not text:
            continue

        # Image captions
        if is_image_caption(paragraph):
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                highlight_alignment(paragraph, report, "Caption under image should be center aligned")
        # Table captions
        elif is_table_caption(paragraph):
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
                highlight_alignment(paragraph, report, "Caption above table should be right aligned")
        # Normal text
        elif paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            highlight_alignment(paragraph, report, "Normal text should be justified")

        if is_image_caption(paragraph) or is_table_caption(paragraph):
            text = paragraph.text.strip()
            # Check caption content
            if is_image_caption(paragraph):
                match = re.match(r"^Рис\.\s*\d+\.\s*(\S+.*)$", text)
                if not match:
                    highlight_alignment(paragraph, report, "Caption must contain text after number")
                    continue
            elif is_table_caption(paragraph):
                match = re.match(r"^Табл\.\s*\d+\.\s*(\S+.*)$", text)
                if not match:
                    highlight_alignment(paragraph, report, "Caption must contain text after number")
                    continue

            # Check caption text format (plain)
            for run in paragraph.runs:
                if run.bold or run.italic or run.underline:
                    highlight_alignment(paragraph, report, "Caption text must be plain (not bold, italic, or underlined)")
                    break

        # Check formatting (including first-line indentation)
        check_paragraph_format(paragraph, report, check_first_line=True)

    # Check paragraphs inside tables
    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    check_paragraph_format(paragraph, report, check_first_line=False)
